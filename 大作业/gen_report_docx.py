# -*- coding: utf-8 -*-
"""Generate formatted Word report from markdown content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), 'single')
            tag.set(qn('w:sz'), '4')
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), '000000')
            tc_borders.append(tag)
    tc_pr.append(tc_borders)


def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=6):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 14, 3: 12}
    set_paragraph_spacing(p, line_spacing=1.5, space_before=12 if level == 1 else 8, space_after=6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True)
    return p


def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_after=6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, line_spacing=1.5, space_after=3)
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.25, space_before=0, space_after=0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else ' ')
        set_run_font(run, name='Consolas', size=10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        set_paragraph_spacing(p, line_spacing=1.25, space_after=0)
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            set_paragraph_spacing(p, line_spacing=1.25, space_after=0)
            run = p.add_run(val)
            set_run_font(run, size=11)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5

    # Title
    add_heading(doc, '鸿蒙南向设备基础课程期末考试报告', level=1)
    add_heading(doc, '试题一：基于 LiteOS-M 任务管理与事件通信的点菜管理系统', level=2)

    # Meta info
    for line in [
        '学号姓名：2023192005 林王儒',
        '课程名称：鸿蒙南向设备开发基础课程',
        '提交日期：2026年6月29日',
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.5, space_after=3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(line), size=12)

    doc.add_paragraph()

    # Section 1
    add_heading(doc, '一、作品背景', level=1)
    add_body(doc, 'OpenHarmony LiteOS-M 内核提供了任务管理（Task）、事件通信（Event）等多种进程间通信（IPC）机制，是嵌入式系统多任务协作的基础。本次期末考试试题一要求学生在 QEMU 仿真环境下，综合运用任务创建与事件读写 API，设计并实现一个支持 Shell 命令交互的点菜管理系统。', indent=True)
    add_body(doc, '该系统模拟餐厅点餐场景：用户在终端通过 menu 命令选择菜品编号下单，后台两个厨师任务分别监听不同菜品对应的事件标志位，收到订单后打印上菜信息。通过本项目，可以深入理解 LiteOS 中事件驱动的多任务协作模型，以及 Shell 用户命令的注册与回调机制。', indent=True)

    # Section 2
    add_heading(doc, '二、功能特点', level=1)
    add_body(doc, '本系统实现了以下核心功能：')
    features = [
        ('事件标志体系', '为 6 种菜品分别定义独立的事件标志位（EVENT_PRODUCT_1 ~ EVENT_PRODUCT_6），通过 ProductInfo 结构体数组将菜品编号、名称、价格、事件位及负责厨师编号绑定在一起。'),
        ('双厨师任务分工', '1 号厨师负责云吞、拌面、炒饭、炒面（编号 1~4）；2 号厨师负责鸡腿饭、猪腿饭（编号 5~6）。两个厨师任务优先级相同（均为 20），启动后进入永久阻塞等待状态。'),
        ('Shell 命令交互', '注册 menu 用户命令，支持 menu <菜品编号> 格式下单；编号非法时给出提示；无参数时打印使用说明及完整菜单。'),
        ('菜单展示', '系统初始化完成后自动打印可读格式的菜品列表，便于用户查阅编号与名称对应关系。'),
        ('统一注册入口', '通过 register_user_test_cmd() 函数完成事件初始化、厨师任务创建、Shell 初始化及命令注册，供 main.c 在系统启动时调用。'),
        ('学号信息输出', '事件系统创建成功后打印学号 2023192005林王儒，满足提交截图要求。'),
    ]
    for i, (title, desc) in enumerate(features, 1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.5, space_after=6)
        p.paragraph_format.first_line_indent = Cm(0.74)
        r1 = p.add_run(f'{i}. {title}  ')
        set_run_font(r1, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2)

    # Section 3
    add_heading(doc, '三、设计思路', level=1)
    add_heading(doc, '3.1 总体架构', level=2)
    add_body(doc, '系统采用生产者—消费者模型：', indent=True)
    add_bullet(doc, '生产者：Shell 命令处理函数 MenuCmdHandler，接收用户输入的菜品编号，调用 LOS_EventWrite 写入对应事件标志位。')
    add_bullet(doc, '消费者：两个厨师任务 Chef1TaskEntry 与 Chef2TaskEntry，分别通过 LOS_EventRead 阻塞等待各自负责的事件掩码，收到事件后查找对应菜品并打印上菜信息。')
    add_body(doc, '事件控制块 g_orderEvent 作为 Shell 命令层与厨师任务层之间的通信桥梁，实现了任务间的解耦与异步通知。', indent=True)

    add_heading(doc, '3.2 事件标志设计', level=2)
    add_body(doc, '采用位掩码方式定义事件标志，每种菜品占用一个独立 bit：', indent=True)
    add_table(doc, ['事件宏', '位值', '对应菜品'], [
        ['EVENT_PRODUCT_1', '1 << 0', '云吞'],
        ['EVENT_PRODUCT_2', '1 << 1', '拌面'],
        ['EVENT_PRODUCT_3', '1 << 2', '炒饭'],
        ['EVENT_PRODUCT_4', '1 << 3', '炒面'],
        ['EVENT_PRODUCT_5', '1 << 4', '鸡腿饭'],
        ['EVENT_PRODUCT_6', '1 << 5', '猪腿饭'],
    ])
    add_body(doc, '厨师任务通过事件掩码过滤只关心的事件：')
    add_bullet(doc, 'CHEF1_EVENT_MASK = EVENT_PRODUCT_1 | ... | EVENT_PRODUCT_4')
    add_bullet(doc, 'CHEF2_EVENT_MASK = EVENT_PRODUCT_5 | EVENT_PRODUCT_6')

    add_heading(doc, '3.3 厨师任务设计', level=2)
    add_body(doc, '两个厨师任务共用 ChefHandleOrder 函数，传入不同的厨师编号、事件掩码及上班提示信息，避免代码重复。任务内部使用 LOS_WAITMODE_OR | LOS_WAITMODE_CLR 模式读取事件：任一匹配事件到达即被唤醒，且读取后自动清除对应标志位，保证每笔订单只处理一次。', indent=True)

    add_heading(doc, '3.4 Shell 命令设计', level=2)
    add_body(doc, '使用 CMD_TYPE_EX 类型注册 menu 命令。该类型会自动屏蔽命令关键字，因此 menu 3 传入回调时 argc=1, argv[0]="3"。命令处理流程为：解析编号 → 查表验证 → 写事件 → 打印下单成功提示。', indent=True)

    # Section 4
    add_heading(doc, '四、软件框架图', level=1)
    framework_lines = [
        '┌─────────────────────────────────────────────────────────────┐',
        '│                      QEMU 仿真终端                           │',
        '│                  用户输入: menu <编号>                        │',
        '└──────────────────────────┬──────────────────────────────────┘',
        '                           │',
        '                           ▼',
        '┌─────────────────────────────────────────────────────────────┐',
        '│              register_user_test_cmd()  [系统入口]             │',
        '│  ┌─────────────────┐  ┌──────────────────────────────────┐  │',
        '│  │ OrderEventInit  │  │ RegisterMenuShellCmd             │  │',
        '│  │ · LOS_EventInit │  │ · OsShellInit                    │  │',
        '│  │ · 创建 Chef1    │  │ · osCmdReg("menu", MenuCmdHandler)│  │',
        '│  │ · 创建 Chef2    │  └──────────────────────────────────┘  │',
        '│  │ · PrintMenu     │                                        │',
        '│  └─────────────────┘                                        │',
        '└──────────────────────────┬──────────────────────────────────┘',
        '                           │',
        '          ┌────────────────┼────────────────┐',
        '          ▼                ▼                ▼',
        '   ┌─────────────┐  ┌─────────────┐  ┌─────────────┐',
        '   │ MenuCmdHandler│ │ Chef1Task   │  │ Chef2Task   │',
        '   │ 解析编号      │  │ 等待掩码1~4 │  │ 等待掩码5~6 │',
        '   │ LOS_EventWrite│  │ 打印上菜    │  │ 打印上菜    │',
        '   └──────┬──────┘  └──────▲──────┘  └──────▲──────┘',
        '          │                │                │',
        '          └────────────────┴────────────────┘',
        '                           │',
        '                           ▼',
        '                  ┌─────────────────┐',
        '                  │  g_orderEvent   │',
        '                  │  (事件控制块)    │',
        '                  └─────────────────┘',
    ]
    add_code_block(doc, framework_lines)

    add_body(doc, '模块关系说明：')
    add_table(doc, ['模块', '职责'], [
        ['register_user_test_cmd', '系统统一初始化入口'],
        ['OrderEventInit', '事件初始化、厨师任务创建、菜单打印'],
        ['RegisterMenuShellCmd', 'Shell 初始化与 menu 命令注册'],
        ['MenuCmdHandler', '用户下单，写入事件标志'],
        ['ChefHandleOrder', '厨师任务通用逻辑，阻塞读事件并上菜'],
        ['ProductInfo 数组', '菜品数据与事件位映射表'],
    ])

    # Section 5
    add_heading(doc, '五、实验步骤', level=1)
    add_heading(doc, '5.1 开发环境准备', level=2)
    for i, step in enumerate([
        '在 Ubuntu 20.04 环境下获取 OpenHarmony 全量源码；',
        '配置交叉编译工具链及 VSCode 开发环境；',
        '将源代码文件放置于对应 Demo 目录，配置 BUILD.gn 编译项；',
        '在启动配置中添加用户命令 Answer1_Demo（或对应题目命令名）。',
    ], 1):
        add_bullet(doc, f'{i}. {step}')

    add_heading(doc, '5.2 代码编写', level=2)
    for i, step in enumerate([
        '定义 ProductInfo 结构体及 6 种菜品数据数组；',
        '实现事件标志宏及厨师事件掩码；',
        '实现 OrderEventInit 完成 LOS_EventInit 与两个厨师任务的 LOS_TaskCreate；',
        '实现 MenuCmdHandler 处理 Shell 下单命令；',
        '实现 register_user_test_cmd 作为统一注册入口，并在 main.c 中调用。',
    ], 1):
        add_bullet(doc, f'{i}. {step}')

    add_heading(doc, '5.3 编译与运行', level=2)
    add_code_block(doc, [
        '# 全量编译（以实际工程路径为准）',
        './build.sh --product-name qemu_small_system_demo --build-target Answer1_Demo',
        '',
        '# 启动 QEMU 仿真，在仿真终端中执行以下命令验证',
        'help          # 确认 menu 命令已注册',
        'menu          # 无参数，打印使用说明和菜单',
        'menu 1        # 下单云吞，1号厨师上菜',
        'menu 5        # 下单鸡腿饭，2号厨师上菜',
        'menu 9        # 非法编号，打印错误提示',
    ])

    add_heading(doc, '5.4 预期运行效果', level=2)
    add_body(doc, '系统启动后依次输出：')
    add_code_block(doc, [
        '<<==>>事件系统创建成功.',
        '<<==>>学号:2023192005林王儒',
        '<<==>>1号厨师上班了.',
        '<<==>>2号厨师上班了.',
        '',
        '*******************点菜系统 :***********************',
        '1,云吞',
        '2,拌面',
        '3,炒饭',
        '4,炒面',
        '5,鸡腿饭',
        '6,猪腿饭',
        '******请输入menu指令,之后空格加对应菜品的数值回车,进行点菜:',
        '<<==>>menu命令注册成功',
    ])
    add_body(doc, '用户执行 menu 2 后：')
    add_code_block(doc, [
        '<<==>>下单成功2.',
        '<<==>>1号厨师送来了:拌面',
    ])

    # Section 6
    add_heading(doc, '六、关键代码说明', level=1)

    add_heading(doc, '6.1 菜品与事件标志绑定', level=2)
    add_code_block(doc, [
        'static ProductInfo g_products[PRODUCT_COUNT] = {',
        '    {1, "云吞",   "12元", EVENT_PRODUCT_1, 1},',
        '    {2, "拌面",   "15元", EVENT_PRODUCT_2, 1},',
        '    {3, "炒饭",   "18元", EVENT_PRODUCT_3, 1},',
        '    {4, "炒面",   "16元", EVENT_PRODUCT_4, 1},',
        '    {5, "鸡腿饭", "20元", EVENT_PRODUCT_5, 2},',
        '    {6, "猪腿饭", "22元", EVENT_PRODUCT_6, 2},',
        '};',
    ])
    add_body(doc, '通过结构体数组统一管理菜品属性，后续扩展或修改只需维护一处数据，符合试题"通过数组将菜品名称与标志位绑定"的要求。', indent=True)

    add_heading(doc, '6.2 厨师任务事件等待与上菜', level=2)
    add_code_block(doc, [
        'static VOID ChefHandleOrder(UINT32 chefNo, UINT32 eventMask, const CHAR *startMsg)',
        '{',
        '    UINT32 readEvent;',
        '    ProductInfo *product;',
        '    printf("%s", startMsg);',
        '    while (1) {',
        '        readEvent = LOS_EventRead(&g_orderEvent, eventMask,',
        '                                  LOS_WAITMODE_OR | LOS_WAITMODE_CLR, LOS_WAIT_FOREVER);',
        '        if (readEvent == 0) continue;',
        '        product = FindProductByEvent(readEvent & eventMask);',
        '        if (product != NULL)',
        '            printf("<<==>>%u号厨师送来了:%s\\n", chefNo, product->name);',
        '    }',
        '}',
    ])
    add_body(doc, '要点说明：')
    add_bullet(doc, 'LOS_WAITMODE_OR：掩码内任一事件置位即唤醒任务；')
    add_bullet(doc, 'LOS_WAITMODE_CLR：读取后自动清除事件标志，防止重复处理；')
    add_bullet(doc, 'LOS_WAIT_FOREVER：无订单时任务永久阻塞，不占用 CPU。')

    add_heading(doc, '6.3 Shell 下单命令处理', level=2)
    add_code_block(doc, [
        'static VOID MenuCmdHandler(INT32 argc, const CHAR *argv[])',
        '{',
        '    if (argc < 1) { PrintMenuUsage(); return; }',
        '    productId = atoi(argv[0]);',
        '    product = FindProductById((UINT32)productId);',
        '    if (product == NULL) { /* 打印编号不合法提示 */ return; }',
        '    ret = LOS_EventWrite(&g_orderEvent, product->eventBit);',
        '    printf("<<==>>下单成功%u.\\n", product->id);',
        '}',
    ])
    add_body(doc, '实现了编号校验、事件写入及结果反馈的完整下单流程。', indent=True)

    add_heading(doc, '6.4 系统统一注册入口', level=2)
    add_code_block(doc, [
        'int register_user_test_cmd(void)',
        '{',
        '    ret = OrderEventInit();',
        '    if (ret != LOS_OK) return (int)ret;',
        '    ret = RegisterMenuShellCmd();',
        '    if (ret != LOS_OK) return (int)ret;',
        '    return 0;',
        '}',
    ])
    add_body(doc, '该函数在 main.c 系统启动阶段被调用，依次完成事件机制初始化、厨师任务创建、Shell 初始化及 menu 命令注册。', indent=True)

    # Section 7
    add_heading(doc, '七、功能完成情况对照', level=1)
    add_table(doc, ['试题要求', '实现情况', '说明'], [
        ['事件标志与菜品绑定', '已实现', '6 种菜品各对应独立事件位，通过数组绑定'],
        ['双厨师任务分工', '已实现', '1号厨师负责1~4号菜，2号厨师负责5~6号菜'],
        ['事件初始化与任务创建', '已实现', 'LOS_EventInit + 两个同优先级厨师任务'],
        ['菜单展示', '已实现', 'PrintMenu 输出编号与名称'],
        ['menu 命令处理', '已实现', '支持下单、非法编号提示、无参数帮助'],
        ['统一注册入口', '已实现', 'register_user_test_cmd'],
        ['学号打印', '已实现', '初始化时输出 2023192005林王儒'],
        ['QEMU 验证', '待截图', '需补充 help 及 menu 命令运行截图'],
    ])

    # Section 8
    add_heading(doc, '八、总结与体会', level=1)
    add_body(doc, '本次大作业通过点菜管理系统的实现，将 LiteOS-M 内核的任务管理与事件通信两个知识点有机结合。主要收获如下：', indent=True)
    summaries = [
        ('事件机制的理解', '事件标志位适合"多对一通知"场景，厨师任务无需轮询，在无订单时处于阻塞态，系统资源利用更高效。'),
        ('任务分工设计', '通过事件掩码将不同职责分配给不同任务，Shell 命令层与业务处理层解耦，结构清晰、易于维护。'),
        ('Shell 命令注册', '掌握了 osCmdReg 与 CMD_TYPE_EX 类型的使用方式，理解了命令参数在回调函数中的传递规则。'),
        ('嵌入式调试经验', '在 QEMU 仿真环境下完成功能验证，熟悉了 OpenHarmony 小系统的编译、烧录与终端交互流程。'),
    ]
    for i, (title, desc) in enumerate(summaries, 1):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.5, space_after=6)
        p.paragraph_format.first_line_indent = Cm(0.74)
        r1 = p.add_run(f'{i}. {title}：')
        set_run_font(r1, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2)
    add_body(doc, '本项目完整实现了基于 LiteOS-M 事件通信的点菜管理系统，达到了试题一的功能要求，为后续学习互斥锁、信号量、消息队列等 IPC 机制打下了实践基础。', indent=True)

    # Appendix
    add_heading(doc, '附录：提交清单', level=1)
    for item in [
        '□ 源代码文件（finalwork.txt / 对应 .c 文件）',
        '□ 头文件及 BUILD.gn 编译配置',
        '□ QEMU 运行截图（含学号打印、help 命令、menu 命令测试）',
        '☑ 大作业实验报告（本文档）',
    ]:
        add_bullet(doc, item)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=1.5, space_before=6)
    r = p.add_run('Git 提交标题格式：大作业提交-2023192005林王儒-试题一')
    set_run_font(r, bold=True)

    out_path = r'D:\南向开发环境\南向作业\大作业\大作业实验报告.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    build_document()
